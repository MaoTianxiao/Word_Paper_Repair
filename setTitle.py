from docx import *
from docx.oxml.ns import qn
from docx.shared import Pt,Inches,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING
 
# # 创建一个已存在的 word 文档的对象
# file = Document('test.docx')
# new_file = Document()
# # print(file)
# # print(file.paragraphs)
# print(file.paragraphs[1].text)


# title_level = 1 # 标题等级
# title_text = file.paragraphs[1].text  # 标题内容
# font_name_ch = u"宋体"  # 标题中文字体
# font_name_en = "Times New Roman"  # 标题英文字体
# font_size = 12 # 标题大小，单位磅
# is_bold = True # 加粗？
# is_italic = False # 斜体？

# space_before_value = 12 # 段前行距
# space_after_value = 12 # 段后行距
# line_space_value = 12 # 行距



# p = new_file.add_heading("",level=2)
# print(p)
# text = p.add_run(title_text)
# text.bold = is_bold
# text.italic = is_italic
# font = text.font
# font.name = font_name_en # 英文字体
# font.size = shared.Pt(font_size)
# # text._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_ch)
# text._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_ch) # 中文字体

# # 设置行距 断距
# paragraph_format=p.paragraph_format
# paragraph_format.space_before=Pt(18)    #上行间距
# paragraph_format.space_after=Pt(12)    #下行间距
# paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE#=Pt(18)  #行距s
# paragraph_format.line_spacing=2.5  #行距s
# paragraph_format.alignment = 0#WD_PARAGRAPH_ALIGNMENT.RIGHT # 居中
# paragraph_format.first_line_indent = 406400 # 首行缩进

# # obj = new_file.add_picture('字号磅数对比.png', width=Inches(1.25))
# # print(type(obj))

# new_file.save('res.docx')


# file = docx.Document("test.docx")
# children = file.element.body.iter()
# child_iters = []
# tags = []
# for child in children:
#     # 通过类型判断目录
#     if child.tag.endswith(('AlternateContent','textbox')):
#         for ci in child.iter():
#             tags.append(ci.tag)
#             if ci.tag.endswith(('main}r', 'main}pPr')):
#                 child_iters.append(ci)
# text = ['']
# for ci in child_iters :
#     if ci.tag.endswith('main}pPr'):
#         text.append('')
#     else:
#         text[-1] += ci.text
#     ci.text = ''
# trans_text = ['***'+t+'***' for t in text]
# print(trans_text)
# i, k = 0, 0
# for ci in child_iters :
#     if ci.tag.endswith('main}pPr'):
#         i += 1
#         k = 0
#     elif k == 0:
#         ci.text = trans_text[i]
#         k = 1
# file.save('12220.docx')
# 提取word目录
    # from lxml import etree
    # children = file.element.body.iter()
    # child_iters = []
    # for child in children:
    #     # 通过类型判断目录
    #     if child.tag.endswith('main}sdt'):
    #         print(etree.tostring(child,encoding='utf-8',pretty_print=True))
    #         # print(child.xml)
    #         print(child.iter() )
    #         print(type(child.iter()))
    #         for ci in child.iter():
    #             if ci.text and ci.text.strip():
    #                 child_iters.append(ci)
    # catalog = [ci.text for ci in child_iters]
    # print(catalog)
file = Document("new.docx")
scs = file.sections
for sc in scs:
    # print(sc.right_margin)
    # sc.right_margin = Cm(3.2)
    # sc.left_margin = Cm(3.2)
    # sc.top_margin = Cm(3.9)
    # sc.bottom_margin = Cm(3.3)
    # sc.header_distance = Cm(1.5)
    # sc.footer_distance = Cm(1.75)
    # print(sc.right_margin)
    print(len(sc.header.paragraphs))
    # print(sc.header.paragraphs[0].text)
    # sc.header.paragraphs[0].text = "1\t\t3"
    sc.header.add_paragraph("1\t2\t3")
print(file.settings.odd_and_even_pages_header_footer)
file.settings.odd_and_even_pages_header_footer = True
print(file.settings.odd_and_even_pages_header_footer)
paras = file.paragraphs
for para in paras:
    for run in para.runs:
        run.bold = True
#     print(para.text)
#     print(para.style.name)
# for table in file.tables:
#     print('12121212121212')
#     print(table.style.name)
#     print( para.paragraph_format.space_before)
#     print( para.paragraph_format.space_after)

#     para.paragraph_format.space_before = Pt(24)
#     para.paragraph_format.space_after = Pt(6)
#     print( para.paragraph_format.space_before)
#     print( para.paragraph_format.space_after)

file.save("1220.docx")
