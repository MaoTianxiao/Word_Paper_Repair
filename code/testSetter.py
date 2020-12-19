from Setter import TextSetter
from docx import Document

setter = TextSetter()
# file = Document('../test.docx')
new_file = Document('../testTitle.docx')
for p in new_file.paragraphs:
    setter.run(p)
new_file.save('test.docx')
