import docx

from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.shape import CT_Inline
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shape import InlineShape

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

file = docx.Document('test.docx')
new_file = docx.Document()
for title1 in iter_block_items(file):
    # title1_style_name = title1.style.name 
    # print(title1_style_name)
    print(type(title1))
    if 'Paragraph' in str(type(title1)):
        temp = new_file.add_paragraph(title1.text)
        temp.element.xml = title1.element.xml
    else:
        temp = new_file.add_table(1,3)
        temp = title1

new_file.save('1219.docx')